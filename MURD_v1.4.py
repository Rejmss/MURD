from PyQt5 import QtCore, QtWidgets, QtGui
from urllib import request
import re
import datetime
from docx import Document
import os
from bs4 import BeautifulSoup


class MainWindow(QtWidgets.QWidget):
    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)

        # Инициализируем файл настроек. Если его нет - создаем пустой, с шаблоном для заполнения
        if os.path.exists(os.path.join(os.getcwd(), 'settings.ini')):
            self.settings = QtCore.QSettings('settings.ini', QtCore.QSettings.IniFormat)
            self.settings.setIniCodec('utf-8')
        else:
            with open('settings.ini', 'w', encoding='utf-8') as file:
                file.write('[Version]\nversion=1.0\n\n[ID]\nid=0\n\n[DefaultLink]\nlink=\n\n[Razdeli]\nrazdeli=,\n\n[Izdanies]\nizdanies=,')
            self.settings = QtCore.QSettings('settings.ini', QtCore.QSettings.IniFormat)
            self.settings.setIniCodec('utf-8')

        # Создаем основной слой.
        self.layout = QtWidgets.QVBoxLayout(self)

        # Создаем табличку и поле для ввода текста статьи
        self.textLabel = QtWidgets.QLabel('Введите текст статьи')
        self.textEditForm = QtWidgets.QTextEdit()
        self.textLayout = QtWidgets.QFormLayout()
        self.textLayout.addRow(QtWidgets.QLabel('Введите текст статьи:'), self.textEditForm)

        # Создаем лейблы сигналов
        self.docsaveLabel = QtWidgets.QLabel('')
        self.sqlsaveLabel = QtWidgets.QLabel('')
        self.filenameUpdateLabel = QtWidgets.QLabel('')

        # Создаем поля ввода
        self.razdelFormLine = QtWidgets.QComboBox()
        self.urlformLine = QtWidgets.QLineEdit()
        self.izdanieFormLine = QtWidgets.QComboBox()
        self.izdanieFormLine.setEditable(True)
        self.izdanieFormLine.setInsertPolicy(6)
        self.izdanieFormLine.setDuplicatesEnabled(False)
        self.idFormLine = QtWidgets.QSpinBox()
        self.idFormLine.setRange(0, 999999)
        self.dateFormLine = QtWidgets.QDateEdit()
        self.dateFormLine.setDisplayFormat('yyyy-MM-dd')
        self.dateFormLine.setDate(datetime.date.today())
        self.titleFormLine = QtWidgets.QLineEdit()
        self.descFormLine = QtWidgets.QTextEdit()
        self.keywordsFormLine = QtWidgets.QLineEdit()
        self.filelinkFormLine = QtWidgets.QLineEdit()
        self.filelinkUpdateLabel = QtWidgets.QLabel('')
        self.duplicateFormLine = QtWidgets.QLineEdit()
        self.filenameLine = QtWidgets.QLineEdit()

        # Создаем кнопки
        self.parsBtn = QtWidgets.QPushButton('Получить данные (Ctrl+F)')
        self.parsBtn.setShortcut('Ctrl+F')
        self.clearBtn = QtWidgets.QPushButton('Очистить все (Ctrl + Q)')
        self.clearBtn.setShortcut('Ctrl+Q')
        self.createSqlBtn = QtWidgets.QPushButton('Дополнить SQL (Ctrl + E)')
        self.createSqlBtn.setShortcut('Ctrl + E')
        self.textSaveBtn = QtWidgets.QPushButton('Сохранить текст статьи (Ctrl + S)')
        self.textSaveBtn.setShortcut('Ctrl + S')
        self.razdelUpdateBtn = QtWidgets.QPushButton('Обновить список разделов')
        self.filenameUpdateBtn = QtWidgets.QPushButton('Обновить имя файла')
        self.filelinkUpdateBtn = QtWidgets.QPushButton('Обновить адрес папки')
        self.articleclearBtn = QtWidgets.QPushButton('Очистить текст статьи')
        self.izdmanupdateBtn = QtWidgets.QPushButton('Обновить список изданий')
        self.htmlkillerBtn = QtWidgets.QPushButton('Убить тэги')
        self.dateincBtn = QtWidgets.QPushButton('')
        self.datedecBtn = QtWidgets.QPushButton('')
        self.dateincBtn.setIcon(QtGui.QIcon('Hopstarter-Button-Button-Next.ico'))
        self.datedecBtn.setIcon(QtGui.QIcon('Hopstarter-Button-Button-Previous.ico'))
        self.dateBtnHbox = QtWidgets.QHBoxLayout()
        self.dateBtnHbox.addWidget(self.datedecBtn)
        self.dateBtnHbox.addWidget(self.dateincBtn)

        # Создаем группу радио-кнопок для выбора режима работы
        self.radbuttonbox = QtWidgets.QGroupBox('Выберите режим работы')
        self.autoRadButton = QtWidgets.QRadioButton('Auto')
        self.manualRadButton = QtWidgets.QRadioButton('Manual')
        self.hbox = QtWidgets.QHBoxLayout()
        self.hbox.addWidget(self.autoRadButton)
        self.hbox.addWidget(self.manualRadButton)
        self.radbuttonbox.setLayout(self.hbox)

        # Создаем блок для контроля очистки текста статьи
        self.safearticleclearBox = QtWidgets.QGroupBox('')
        self.safecheck = QtWidgets.QCheckBox('Подтверждение очистки')
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.articleclearBtn)
        self.vbox.addWidget(self.safecheck)
        self.safearticleclearBox.setLayout(self.vbox)
        self.safecheck.setChecked(True)

        # Создаем блок выбора пути сохранения файлов
        self.savecheck = QtWidgets.QCheckBox('Групповая обработка')
        self.savecheck.setChecked(False)

        # Добавляем поля ввода на соответсвующий слой
        self.formBoxLayout = QtWidgets.QFormLayout()
        self.formBoxLayout.addRow(QtWidgets.QLabel('SQL ID:'), self.idFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('DATE:'), self.dateFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel(''), self.dateBtnHbox)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Раздел:'), self.razdelFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Издание:'), self.izdanieFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Заголовок:'), self.titleFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Имя файла:'), self.filenameLine)
        self.formBoxLayout.addRow(self.filenameUpdateLabel, self.filenameUpdateBtn)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Описание:'), self.descFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Ключевые слова:'), self.keywordsFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Ссылка:'), self.urlformLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Путь к файлу:'), self.filelinkFormLine)
        self.formBoxLayout.addRow(QtWidgets.QLabel(''), self.savecheck)
        self.formBoxLayout.addRow(self.filelinkUpdateLabel, self.filelinkUpdateBtn)
        self.formBoxLayout.addRow(QtWidgets.QLabel('Duplicate:'), self.duplicateFormLine)

        # Генерируем вкладки
        self.tabs = QtWidgets.QTabWidget()
        self.tabInp = QtWidgets.QWidget()
        self.tabOutp = QtWidgets.QWidget()
        self.tabs.addTab(self.tabInp, 'Внесение данных')
        self.tabs.addTab(self.tabOutp, 'Сформировать выгрузку')

        # Генерируем создержимое первой вкладки
        self.tabInp.layout = QtWidgets.QGridLayout(self)
        self.tabInp.layout.addLayout(self.formBoxLayout, 0, 0, 10, 2)
        self.tabInp.layout.addWidget(self.parsBtn, 10, 0, 1, 1)
        self.tabInp.layout.addWidget(self.textLabel, 0, 2, 1, 2, alignment=QtCore.Qt.AlignCenter)
        self.tabInp.layout.addWidget(self.textEditForm, 1, 2, 17, 2)
        self.tabInp.layout.addWidget(self.textSaveBtn, 12, 0, 1, 1)
        self.tabInp.layout.addWidget(self.docsaveLabel, 12, 1, 1, 1)
        self.tabInp.layout.addWidget(self.createSqlBtn, 13, 0, 1, 1)
        self.tabInp.layout.addWidget(self.sqlsaveLabel, 13, 1, 1, 1)
        self.tabInp.layout.addWidget(self.clearBtn, 14, 0, 1, 1)
        self.tabInp.layout.addWidget(self.razdelUpdateBtn, 15, 0, 1, 1)
        self.tabInp.layout.addWidget(self.safearticleclearBox, 14, 1, 2, 1)
        self.tabInp.layout.addWidget(self.izdmanupdateBtn, 16, 0, 1, 1)
        self.tabInp.layout.addWidget(self.htmlkillerBtn, 16, 1, 1, 1)
        self.tabInp.layout.addWidget(self.radbuttonbox, 17, 0, 1, 1)
        self.autoRadButton.setChecked(True)

        self.tabInp.setLayout(self.tabInp.layout)
        # Добавляем виджет вкладок к нашему окну
        self.layout.addWidget(self.tabs)
        self.setLayout(self.layout)

        # Валидатор корректности заполнения имени файла
        def filenameupdate():
            stoppattern = '~|#|%|&|\*|{|}|\\|:|<|>|\?|/|\+|\||\"|»|«'
            newfilename = f'{self.idFormLine.text()} - {self.titleFormLine.text()}'

            for i in stoppattern:
                newfilename = newfilename.replace(i, '')

            self.filenameLine.setText(newfilename)
            self.filenameUpdateLabel.setText('Обновлено')
            self.filenameUpdateLabel.setStyleSheet(('color: green'))

            if re.search(stoppattern, self.filenameLine.text()):
                self.filenameLine.setStyleSheet('background-color: red')
            else:
                self.filenameLine.setStyleSheet('background-color: green')
            self.parsBtn.setDisabled(True)

        # Формирование и сохранение файла в формате .docx
        def docxdocsave():
            document = Document()

            if self.filenameLine.text() == '':
                self.filenameLine.setStyleSheet('background-color: red')
                self.docsaveLabel.setText('Некорректное имя файла')
                self.docsaveLabel.setStyleSheet('color: red')
            if self.textEditForm.toPlainText() == '':
                self.docsaveLabel.setText('Отсутствует текст статьи')
                self.docsaveLabel.setStyleSheet('color: red')
            else:
                document.add_heading(self.titleFormLine.text(), level=2)
                document.add_paragraph('\n' + self.dateFormLine.text())
                document.add_paragraph('\n' + self.razdelFormLine.currentText())
                document.add_paragraph('\n' + self.keywordsFormLine.text())
                document.add_paragraph('\n' + self.urlformLine.text())
                document.add_paragraph('\n' + self.textEditForm.toPlainText())

                if not self.savecheck.isChecked():
                    filelink_update()

                filelink = self.filelinkFormLine.text()

                filename = self.filenameLine.text() + '.docx'

                fullfilelink = os.path.join(filelink, filename)
                document.save(fullfilelink)
                endwork = f'{self.dateFormLine.text()}.docx'

                os.listdir(filelink)
                if filename in os.listdir(filelink):
                    self.docsaveLabel.setText('Сохранение выполнено')
                    self.docsaveLabel.setStyleSheet('color: green')
                if endwork not in os.listdir(filelink):
                    with open(os.path.join(filelink, endwork), 'w') as file:
                        pass
                    self.docsaveLabel.setText('Сохранение выполнено || Итоговый файл создан')

        # Функция-парсер
        def pars():

            if self.urlformLine.text() == '':
                self.urlformLine.setStyleSheet('background-color: red')
            else:
                self.urlformLine.setStyleSheet('background-color: white')

                urlpattern = '/$'
                if re.search(urlpattern, self.urlformLine.text()):
                    realurl = self.urlformLine.text().replace(urlpattern, '')
                    self.urlformLine.setText(realurl)

                url = self.urlformLine.text()

                hdr = {
                    'user-agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.103 YaBrowser/19.4.2.702 Yowser/2.5 Safari/537.36'
                }

                req = request.Request(
                    url,
                    data=None,
                    headers=hdr
                )

                resp = BeautifulSoup(request.urlopen(req), 'lxml')

                title = resp.title
                if title:
                    title = title.string
                    title = title.replace('\n', '')
                    if title.isupper():
                        title = title.lower

                description = resp.find('meta', attrs={'name': 'description'})
                if description:
                    description = description.get('content')

                keywords = resp.find('meta', attrs={'name': 'keywords'})
                if keywords:
                    keywords = keywords.get('content')

                article = resp.findAll('p')
                if article:
                    mylist = ['</p>, <p>', '<p>', '[', ']', '</p>']
                    for i in mylist:
                        article = str(article).replace(i, ' ')
                else:
                    article = ':-('

            self.parsBtn.setDisabled(True)

            pars_rezult(title, description, keywords, article)

        # Заполнение полей для парсера
        def pars_rezult(title, description, keywords, article):
            self.titleFormLine.setText(title)
            self.descFormLine.setText(description)
            self.keywordsFormLine.setText(keywords)
            self.textEditForm.setText(article)
            self.duplicateFormLine.setText('NULL')

            izdanie = self.urlformLine.text()[
                      :self.urlformLine.text().find('/', self.urlformLine.text().find('//') + 2)]
            self.izdanieFormLine.setCurrentText(izdanie)
            izdanies_checking()
            filenameupdate()

        # Функция генерации sql-инжектора
        def sqlfileupdate():
            if not self.savecheck.isChecked():
                filelink_update()

            filelink = self.filelinkFormLine.text()

            filename = 'sql-listing.txt'
            fullfilelink = os.path.join(filelink, filename)

            if self.duplicateFormLine.text() == '':
                self.duplicateFormLine.setText('NULL')

            if self.idFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.dateFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.titleFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.descFormLine.toPlainText() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.keywordsFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.urlformLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.filelinkFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            elif self.duplicateFormLine.text() == '':
                self.sqlsaveLabel.setText('Необходимо заполнить все поля!')
                self.sqlsaveLabel.setStyleSheet('color: red')
            else:
                with open(fullfilelink, 'a') as file:
                    file.write(
                        '\n' + "INSERT INTO `work_reputation`.`main_table`(`id`, `date`, `razdel`, `izdanie`, `title`,"
                               " `description`, `keyword`, `hyperlink`, `file_link`, `duplicate`) "
                               "VALUES ('{id}', '{date}', '{razdel}', '{izdanie}', '{title}', '{description}', '{keyword}', "
                               "'{hyperlink}', '{file_link}', '{duplicate}');".format(id=self.idFormLine.text(),
                                                                                      date=self.dateFormLine.text(),
                                                                                      razdel=self.razdelFormLine.currentText(),
                                                                                      izdanie=self.izdanieFormLine.currentText(),
                                                                                      title=self.titleFormLine.text(),
                                                                                      description=self.descFormLine.toPlainText(),
                                                                                      keyword=self.keywordsFormLine.text(),
                                                                                      hyperlink=self.urlformLine.text(),
                                                                                      file_link=self.filelinkFormLine.text(),
                                                                                      duplicate=self.duplicateFormLine.text()))
                    self.sqlsaveLabel.setText('Сохранение выполнено')
                    self.sqlsaveLabel.setStyleSheet('color: green')

        # Функция очистки полей
        def clearfields():
            self.izdanieFormLine.setCurrentIndex(0)
            self.titleFormLine.clear()
            self.descFormLine.clear()
            self.keywordsFormLine.clear()
            self.textEditForm.clear()
            self.filenameLine.clear()
            if self.sqlsaveLabel.text() == 'Сохранение выполнено' \
                    and (self.docsaveLabel.text() == 'Сохранение выполнено'
                         or self.docsaveLabel.text() == 'Сохранение выполнено || Итоговый файл создан'):
                self.idFormLine.stepUp()
            self.settings.setValue("ID/id", self.idFormLine.text())
            self.razdelFormLine.setCurrentIndex(0)
            self.urlformLine.clear()
            self.duplicateFormLine.clear()
            self.docsaveLabel.setText(' ')
            self.sqlsaveLabel.setText(' ')
            self.filenameUpdateLabel.setText(' ')
            self.filenameLine.setStyleSheet('background-color: white')
            self.parsBtn.setEnabled(True)
            self.safecheck.setChecked(True)

        # Ручная очистка поля статьи (и только его!)
        def articleclear():
            if self.safecheck.isChecked():
                self.textEditForm.clear()
                self.safecheck.setChecked(False)

        # Обновление списка разделов
        def razdelupdate():
            self.settings.sync()
            self.razdelFormLine.clear()
            for i in self.settings.value("Razdeli/razdeli"):
                if i == '___________________':
                    self.razdelFormLine.insertSeparator(len(self.razdelFormLine))
                    continue
                self.razdelFormLine.addItem(i)
                self.razdelFormLine.setMaxVisibleItems(len(self.razdelFormLine))

        # Инициализатор данных о изданиях
        def izdaniesinitiate():
            self.settings.sync()
            for i in self.settings.value("Izdanies/izdanies"):
                self.izdanieFormLine.addItem(i)
                self.izdanieFormLine.setMaxVisibleItems(8)

        # Проверка изданий на соответствие файлу настроек
        def izdanies_checking():
            izdanie = self.izdanieFormLine.currentText()
            ilist = self.settings.value("Izdanies/izdanies")
            if izdanie in ilist:
                pass
            else:
                ilist.append(izdanie)
                self.settings.setValue("Izdanies/izdanies", ilist)
            self.izdanieFormLine.clear()
            izdaniesinitiate()
            self.izdanieFormLine.setCurrentText(izdanie)

        # Ручное обновление изданий.
        def izdanies_manual_update():
            if self.izdanieFormLine.currentText() != '' or self.izdanieFormLine.currentText() != ' ':
                izdanies_checking()
            else:
                self.izdanieFormLine.clear()
                izdaniesinitiate()
            ilist = self.settings.value("Izdanies/izdanies")
            if '' in ilist:
                ilist.remove('')
                self.settings.setValue("Izdanies/izdanies", ilist)

        # Ручная установка пути для папки
        def filelink_manual_set():
            dirname = QtWidgets.QFileDialog.getExistingDirectory(directory=f'{self.settings.value("DefaultLink/link")}'
            f'{self.dateFormLine.text()}')
            self.filelinkFormLine.setText(dirname)

        # Автоматическая проверка пути до рабочей директории соответственно файлу настроек
        def filelink_update():
            filelink = f'{self.settings.value("DefaultLink/link")}{self.dateFormLine.text()}'

            drive, b = os.path.splitdrive(filelink)

            if os.path.exists(drive):
                if not os.path.exists(filelink):
                    os.makedirs(filelink)
                self.filelinkFormLine.setText(filelink)
            else:
                filelink = os.path.join(os.getcwd(), self.dateFormLine.text())
                if not os.path.exists(filelink):
                    os.makedirs(filelink)
                self.filelinkFormLine.setText(filelink)

        # Обновление ID
        def id_update():
            self.idFormLine.setValue(int(self.settings.value("ID/id")))

        # Смена режима работы
        def work_mod_check():
            if self.manualRadButton.isChecked():
                self.parsBtn.setDisabled(True)
            else:
                self.parsBtn.setDisabled(False)

        # Очитска основных тэгов из статьи
        def htmlkiller():
            text = self.textEditForm.toPlainText()
            mylist = ['<p>', '[', ']', '</p>, <p>', '</p>', '<br>', '<br/>', '<br />']
            for i in mylist:
                text = str(text).replace(i, '')
            self.textEditForm.setText(text)

        def dateinc():
            self.dateFormLine.setCurrentSectionIndex(2)
            self.dateFormLine.stepUp()

        def datedec():
            self.dateFormLine.setCurrentSectionIndex(2)
            self.dateFormLine.stepDown()

        # Блок инициализирующих функций
        id_update()
        razdelupdate()
        filelink_update()
        izdaniesinitiate()

        # Подключение помощника автозаполнения изданий
        self.izdaniescompleter = QtWidgets.QCompleter(self.settings.value("Izdanies/izdanies"), self.izdanieFormLine)
        self.izdanieFormLine.setCompleter(self.izdaniescompleter)

        # Подключение функций к кнопкам
        self.textSaveBtn.clicked.connect(docxdocsave)
        self.parsBtn.clicked.connect(pars)
        self.createSqlBtn.clicked.connect(sqlfileupdate)
        self.clearBtn.clicked.connect(clearfields)
        self.razdelUpdateBtn.clicked.connect(razdelupdate)
        self.filenameUpdateBtn.clicked.connect(filenameupdate)
        self.filelinkUpdateBtn.clicked.connect(filelink_manual_set)
        self.articleclearBtn.clicked.connect(articleclear)
        self.izdmanupdateBtn.clicked.connect(izdanies_manual_update)
        self.manualRadButton.clicked.connect(work_mod_check)
        self.autoRadButton.clicked.connect(work_mod_check)
        self.htmlkillerBtn.clicked.connect(htmlkiller)
        self.dateincBtn.clicked.connect(dateinc)
        self.datedecBtn.clicked.connect(datedec)


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.setWindowTitle('MURD v1.4')
    window.resize(1200, 700)
    window.show()
    sys.exit(app.exec_())
